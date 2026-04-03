"""Generate ES_Parachute.pdf — NVC Executive Summary with strict formatting."""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page setup: 1-inch margins ──
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(12)
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_after = Pt(0)
pf.space_before = Pt(0)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(14) if level == 1 else Pt(13)
        run.font.name = 'Calibri'
        run.font.color.rgb = None  # inherit
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.line_spacing = 1.5
    return h


def add_para(text, bold_prefix=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = space_after
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(12)
        run_b.font.name = 'Calibri'
        run_rest = p.add_run(text)
        run_rest.font.size = Pt(12)
        run_rest.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = 'Calibri'
    return p


def add_bullet(text, bold_prefix=None, space_after=Pt(4)):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = space_after
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(12)
        run_b.font.name = 'Calibri'
        run_rest = p.add_run(text)
        run_rest.font.size = Pt(12)
        run_rest.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = 'Calibri'
    return p


# ═══════════════════════════════════════════════
# COVER / HEADER
# ═══════════════════════════════════════════════
title = doc.add_heading('Open Parachute PBC — Executive Summary', level=0)
for run in title.runs:
    run.font.size = Pt(18)
    run.font.name = 'Calibri'
title.paragraph_format.space_after = Pt(8)

add_para('Open Parachute PBC (Colorado Public Benefit Corporation)', bold_prefix='Startup: ', space_after=Pt(2))
p = add_para('aaron@parachute.computer', bold_prefix='Founder: Aaron Gabriel Neyer — ', space_after=Pt(2))
add_para('Jon Bo (Daily Co-lead), Lucian Hymer (Computer Co-lead), Marvin Melzer (Hardware), Neil Yarnal (Brand & Design)', bold_prefix='Team: ', space_after=Pt(2))
add_para('Boulder, Colorado', bold_prefix='Location: ', space_after=Pt(8))


# ═══════════════════════════════════════════════
# OPPORTUNITY SUMMARY
# ═══════════════════════════════════════════════
add_heading('Opportunity Summary', level=1)

add_para(
    'Personal agentic computing is the defining shift in how people interact with technology. '
    'In the last six months, tools like OpenClaw (300,000+ GitHub stars in three months), Claude Cowork, '
    'ZoComputer, and Perplexity Computer have validated that people want an AI that works for them \u2014 '
    'not just answering questions but taking action, managing tools, and learning who they are over time. '
    'Over 100 million people already pay $20+/month for AI \u2014 including 50M+ ChatGPT Pro subscribers. '
    'The agentic AI market is projected to exceed $50B by 2030 at a 44% CAGR, and OpenAI alone projects '
    '$200B in annual revenue by that year.'
)

add_para(
    'But nearly every player is building for the power user \u2014 the person already paying $20\u2013200/month '
    'and comfortable handing an AI agent full access to their digital life. That\u2019s roughly 5% of the '
    'addressable market. The other 95% \u2014 artists, small business owners, everyday people who know AI could '
    'help but don\u2019t know where to start \u2014 have no accessible entry point into this future. Parachute bridges '
    'that gap with two products that form one journey: Parachute Daily, a voice-first journaling app that gently '
    'introduces AI into daily life, and Parachute Computer, a full open-source agentic platform. Daily builds '
    'the context that makes Computer powerful, creating a natural upgrade path and a compounding moat that no '
    'competitor can shortcut.'
)


# ═══════════════════════════════════════════════
# PRODUCT OR SERVICE
# ═══════════════════════════════════════════════
add_heading('Product or Service', level=1)

add_para(
    ' is a voice-first journal. Users speak into a wearable pendant or their phone \u2014 '
    'on a walk, in the car, wherever thinking happens. Entries are transcribed (offline-capable via on-device '
    'models), organized, and enhanced with AI-powered daily reflections, pattern recognition, and weekly synthesis. '
    'It works fully offline as a simple journal for free; cloud transcription unlocks at $5/month and AI features '
    'at $10/month.',
    bold_prefix='Parachute Daily'
)

add_para(
    ' is a full personal agentic computing platform. It includes a knowledge graph (Brain) '
    'that connects journals, conversations, and structured data into a unified model of the user\u2019s thinking. '
    'It supports multi-agent teams, trust-tiered execution (from sandboxed to full system access), and connectors '
    'to Telegram, Discord, and other messaging platforms. Available as a hosted service at $40/month or fully '
    'self-hosted for free.',
    bold_prefix='Parachute Computer'
)

add_para(
    'The critical insight is that context compounds. Every journal entry, every conversation, every voice note '
    'builds a richer knowledge graph. After months of Daily use, a user\u2019s system already understands how they '
    'think, what they care about, and what they\u2019re working on. When they\u2019re ready for more, they upgrade to '
    'Parachute Computer and their brain comes with them. This compounding context is both the core user value '
    'and the primary switching cost \u2014 no competitor can replicate months of accumulated personal context.'
)

add_para(
    ' Working Python/FastAPI server, Flutter app (macOS, Android, web), graph-native memory '
    'infrastructure, local voice transcription via Sherpa-ONNX, multi-agent system, and three bot connectors '
    '(Telegram, Discord, Matrix). Functional pendant prototype with custom Parachute enclosure. Daily beta '
    'launching this month, with a polished production launch targeted for June 2026. OpenParachute PBC '
    'incorporated in Colorado.',
    bold_prefix='Current state of development:'
)


# ═══════════════════════════════════════════════
# COMPETITIVE DIFFERENTIATION
# ═══════════════════════════════════════════════
add_heading('Competitive Differentiation', level=1)

add_para(
    'The personal AI space is crowded and accelerating. Key players include OpenClaw (open-source agentic platform), '
    'TwinMind ($5.7M raised at $60M valuation), Mem.ai ($28.6M raised), ZoComputer, Perplexity Computer, and Manus. '
    'Parachute differentiates on three axes:'
)

add_bullet(
    ' Competitors target power users who are already bought into AI. Parachute Daily gives everyday people a simple, '
    'voice-first entry point that requires no technical sophistication \u2014 just talk. This opens the mass market that '
    'every other player is ignoring.',
    bold_prefix='The bridge to the 95%. '
)

add_bullet(
    ' An AI works best when it knows everything about you. People will only share that depth of context with a system '
    'they trust. Parachute is fully open source (AGPL-3.0) and local-first \u2014 data lives on the user\u2019s device, '
    'portable and exportable. As a Public Benefit Corporation, our legal structure mandates that we consider the '
    'interests of our users, not just our shareholders.',
    bold_prefix='Open source as trust.'
)

add_bullet(
    ' Software can be cloned in a day. Compounding personal context cannot. Every day a user journals, reflects, '
    'and converses with their AI, the switching cost grows \u2014 not through lock-in, but through genuine accumulated '
    'value that no competitor can replicate.',
    bold_prefix='Context compounds as the real moat.'
)


# ═══════════════════════════════════════════════
# MARKET AND CUSTOMER ANALYSIS
# ═══════════════════════════════════════════════
add_heading('Market and Customer Analysis', level=1)

add_para(
    'The agentic AI market is projected to exceed $50B by 2030 at a 44% CAGR (MarketsandMarkets, Capgemini/Statista). '
    'OpenAI alone projects $200B in annual revenue by 2030. Adjacent comparables demonstrate strong investor interest '
    'and viable business models: TwinMind ($5.7M raised at $60M valuation), Obsidian (~$25M ARR as a note-taking tool), '
    'Day One (~$4.8M ARR as a journaling app), and Mem.ai ($28.6M raised for AI-powered knowledge management).'
)

add_para('We target two customer segments through one funnel:')

add_bullet(
    ' People who aren\u2019t yet bought into AI but will journal, capture thoughts, and gradually experience AI\u2019s '
    'value in their lives. This is the 95% \u2014 the artist who wants to organize creative ideas, the small business '
    'owner tracking their days, the parent who wants a better way to remember and reflect. Entry at free, cloud '
    'transcription at $5/month, AI features at $10/month.',
    bold_prefix='Daily users (mass market):'
)

add_bullet(
    ' Builders and professionals who want a full agentic AI platform they own and trust. $40/month hosted or free '
    'self-hosted. These users also create tools, integrations, and workflows that benefit the broader ecosystem.',
    bold_prefix='Computer users (power users):'
)

add_para(
    ' 300+ community members in our Boulder ecosystem ready to onboard. 13 builders completed our first Learn Vibe '
    'Build AI learning cohort. Active users in private beta providing ongoing feedback.',
    bold_prefix='Validation:'
)


# ═══════════════════════════════════════════════
# INTELLECTUAL PROPERTY
# ═══════════════════════════════════════════════
add_heading('Intellectual Property', level=1)

add_para(
    'Parachute is open source under the AGPL-3.0 license. This is a deliberate strategic choice: AGPL requires '
    'that anyone who runs a modified version of Parachute as a service must share their changes, protecting against '
    'proprietary forks competing against us. Our defensible advantages are the compounding user context (which lives '
    'with each user), the knowledge graph architecture, the community ecosystem, and the trust earned by building '
    'in the open. We believe open source is a competitive advantage \u2014 it builds the trust necessary for people '
    'to share their thinking with an AI system.'
)


# ═══════════════════════════════════════════════
# MANAGEMENT TEAM
# ═══════════════════════════════════════════════
add_heading('Management Team', level=1)

add_bullet(
    ' Product vision and system architecture. MA in Ecopsychology, MS in Creative Technology & Design '
    '(CU Boulder ATLAS). Founding engineer at two startups. Former Google. 10+ years full stack development. '
    'Boulder Human Relations Commission Chair. Founder of Woven Web (501(c)(3) nonprofit).',
    bold_prefix='Aaron Gabriel Neyer (Founder) \u2014'
)
add_bullet(' Daily co-lead. Founding engineer at multiple startups. Leading product direction for Parachute Daily.', bold_prefix='Jon Bo \u2014')
add_bullet(' Computer co-lead. Founding engineer at multiple startups. Leading architecture for the full agentic platform.', bold_prefix='Lucian Hymer \u2014')
add_bullet(' Hardware lead. Developing the wearable pendant prototype. Experienced hardware designer and engineer.', bold_prefix='Marvin Melzer \u2014')
add_bullet(' Brand and design.', bold_prefix='Neil Yarnal \u2014')

add_para('An additional 3\u20134 experienced builders are available for hire as funding increases, enabling the team to scale from a core of 4\u20135 to 9\u201310.')


# ═══════════════════════════════════════════════
# FINANCIAL PROJECTIONS
# ═══════════════════════════════════════════════
add_heading('Financial Projections', level=1)

# Financial table
headers = ['', '2026', '2027', '2028']
data = [
    ['Free + sync users', '5,000', '50,000', '250,000'],
    ['Paid subscribers ($2\u201340/mo)', '500', '5,000', '25,000'],
    ['Avg rev / paid subscriber', '~$7/mo', '~$9/mo', '~$12/mo'],
    ['ARR', '~$42K', '~$540K', '~$3.6M'],
    ['Team costs', '~$150K', '~$400K', '~$800K'],
    ['Infra + COGS', '~$15K', '~$130K', '~$500K'],
    ['Total opex', '~$165K', '~$530K', '~$1.3M'],
]

table = doc.add_table(rows=len(data) + 1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# Header row
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'

# Data rows
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        cell = table.cell(r + 1, c)
        cell.text = val
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                if r == 3 and c > 0:  # ARR row highlight
                    run.bold = True

# Spacing after table
doc.add_paragraph()

add_para(
    ' Tiered subscriptions \u2014 Free (offline journal), $2/mo (cloud sync), $5/mo (cloud transcription + cleanup), '
    '$10/mo (AI reflections, synthesis, pattern surfacing), $40/mo (hosted Parachute Computer). The free tier is '
    'fully offline with no hosting cost \u2014 no subsidizing free users. Infrastructure costs include cloud '
    'transcription (~$1.50/1000 min) and AI inference (cost-efficient models for reflections). As model costs '
    'continue to decline, margins improve further.',
    bold_prefix='Revenue model:'
)

add_para(
    ' Year two approaches breakeven (~$540K ARR vs ~$530K total opex). Year three is clearly profitable '
    '(~$3.6M ARR vs ~$1.3M total opex) as the subscriber mix shifts toward higher tiers and infrastructure '
    'costs scale sub-linearly. Average revenue per subscriber increases over time as users\u2019 context compounds '
    'and they naturally upgrade \u2014 the product gets more valuable the longer you use it, which drives organic '
    'upselling.',
    bold_prefix='Path to profitability:'
)

add_para(
    ' $0. The entire product has been self-funded with no outside investment.',
    bold_prefix='Funding to date:'
)


# ═══════════════════════════════════════════════
# INVESTMENT
# ═══════════════════════════════════════════════
add_heading('Investment', level=1)

add_para(
    'Raising $300,000 via a SAFE note (YC standard) at a $5,000,000 valuation cap. This is deliberately priced as '
    'early-believer terms \u2014 TwinMind raised at a $60M valuation with 30,000 users; we are raising before public launch.'
)

add_para(
    ' Bring the core team (founder + two co-leads) full-time through 2026, fund infrastructure and hosting costs, '
    'and execute a polished production launch by June 2026. The goal is revenue immediately upon launch, with '
    'sufficient growth to raise a subsequent round by early 2027 at a significantly higher valuation, scaling the '
    'team from 4\u20135 to 9\u201310.',
    bold_prefix='Use of funds:'
)

add_para(
    'NVC prize funds would accelerate this timeline \u2014 enabling faster team ramp-up and broader beta distribution.',
    bold_prefix=''
)


# ── Save ──
docx_path = '/home/sandbox/parachute-computer/website/nvc/ES_Parachute.docx'
doc.save(docx_path)
print(f'Saved Word doc to {docx_path}')
